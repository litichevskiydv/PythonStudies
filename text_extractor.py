import os
import re
import pathlib
from docx import Document
from win32com import client

text_pattern = re.compile('[a-zA-Zа-яА-Я0-9]+')


def extract_text_from_docx(file_path):
    paragraphs = []
    document = Document(file_path)

    for paragraph in document.paragraphs:
        if text_pattern.match(paragraph.text) is not None:
            paragraphs.append(paragraph.text)

    paragraphs.append('')
    for table in document.tables:
        for row in table.rows:
            paragraphs.append('\t'.join([cell.text for cell in row.cells]))

    return paragraphs


def extract_text_from_doc(file_path):
    doc_full_path = os.path.join(pathlib.Path().resolve(), file_path)
    docx_full_path = os.path.splitext(doc_full_path)[0] + '.' + 'docx'

    word = client.Dispatch('Word.Application')
    document = word.Documents.Open(doc_full_path)
    document.SaveAs(docx_full_path, 12)
    document.Close()
    word.Quit()

    return extract_text_from_docx(docx_full_path)


def main():
    extracted_text = '\n'.join(extract_text_from_doc('...'))
    print(extracted_text)


if __name__ == '__main__':
    main()
